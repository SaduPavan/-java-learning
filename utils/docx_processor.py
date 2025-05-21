from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def load_docx(file_path):
    return Document(file_path)


def enhance_paragraph_style(paragraph):
    # Example: Make headings blue and bold
    if paragraph.style.name.startswith('Heading'):
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.color.rgb = RGBColor(0, 51, 153)  # Professional blue
        run.bold = True
    # Center align titles
    if paragraph.style.name == 'Title':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def enhance_table_style(table):
    # Apply a professional style to tables
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text


def process_docx(input_path, output_path):
    doc = load_docx(input_path)
    for paragraph in doc.paragraphs:
        enhance_paragraph_style(paragraph)
    for table in doc.tables:
        enhance_table_style(table)
    doc.save(output_path) 